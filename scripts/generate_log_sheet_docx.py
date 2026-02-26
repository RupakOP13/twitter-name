from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _set_paragraph_text(cell, text: str, *, bold: bool = False, size_pt: float = 11, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)


def _set_cell_border(cell, **kwargs):
    """Set cell`s border.

    Usage:
        _set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "000000"},
            bottom={...},
            left={...},
            right={...},
        )

    sz is in eighths of a point (Word units). 4 ~= 0.5pt, 12 ~= 1.5pt.
    """

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("left", "top", "right", "bottom"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcBorders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def generate(out_path: Path) -> None:
    doc = Document()

    # Page setup: aim to fit the entire table on one page.
    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4
    section.page_height = Inches(11.69)  # A4
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)

    # Table structure (matches image):
    # Columns: Date | Time: From | To | Total Hours | Details of work done | Signature of Authorised Person | Signature of student
    cols = 7
    blank_rows = 12
    rows = 1 + 2 + blank_rows + 1  # title + (header main + header sub) + blank + total

    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = False

    # Column widths tuned to A4 with margins
    col_widths = [
        Inches(0.80),  # Date
        Inches(0.60),  # Time From
        Inches(0.60),  # Time To
        Inches(0.70),  # Total Hours
        Inches(2.97),  # Details
        Inches(0.90),  # Signature (Authorised)
        Inches(0.90),  # Signature (Student)
    ]

    for i, w in enumerate(col_widths):
        for r in range(rows):
            table.cell(r, i).width = w

    # Default borders (thin) + vertical centering
    thin = {"sz": 4, "val": "single", "color": "000000"}
    for r in range(rows):
        row = table.rows[r]
        # Keep row heights compact but readable
        row.height = Inches(0.42 if r == 0 else 0.45)
        for c in range(cols):
            cell = row.cells[c]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_border(cell, top=thin, bottom=thin, left=thin, right=thin)
            # Make paragraphs compact
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1

    # Title row (merge across all columns)
    title_cell = table.cell(0, 0)
    for c in range(1, cols):
        title_cell = title_cell.merge(table.cell(0, c))
    _set_paragraph_text(table.cell(0, 0), "LOG SHEET OF WORK PERFORMED DURING INTERNSHIP", bold=True, size_pt=13)

    # Header rows
    header_main_r = 1
    header_sub_r = 2

    # Merge vertically for multi-row headers except Time.
    table.cell(header_main_r, 0).merge(table.cell(header_sub_r, 0))  # Date
    table.cell(header_main_r, 3).merge(table.cell(header_sub_r, 3))  # Total Hours
    table.cell(header_main_r, 4).merge(table.cell(header_sub_r, 4))  # Details
    table.cell(header_main_r, 5).merge(table.cell(header_sub_r, 5))  # Signature (Auth)
    table.cell(header_main_r, 6).merge(table.cell(header_sub_r, 6))  # Signature (Student)

    # Time header merges across From/To
    table.cell(header_main_r, 1).merge(table.cell(header_main_r, 2))

    _set_paragraph_text(table.cell(header_main_r, 0), "Date", bold=True, size_pt=11)
    _set_paragraph_text(table.cell(header_main_r, 1), "Time", bold=True, size_pt=11)
    _set_paragraph_text(table.cell(header_sub_r, 1), "From", bold=True, size_pt=11)
    _set_paragraph_text(table.cell(header_sub_r, 2), "To", bold=True, size_pt=11)
    _set_paragraph_text(table.cell(header_main_r, 3), "Total\nHours", bold=True, size_pt=11)
    _set_paragraph_text(table.cell(header_main_r, 4), "Details of work done", bold=True, size_pt=11)
    _set_paragraph_text(table.cell(header_main_r, 5), "Signature\nof\nAuthorised\nPerson", bold=True, size_pt=10)
    _set_paragraph_text(table.cell(header_main_r, 6), "Signature\nof student", bold=True, size_pt=10)

    # Blank data rows (leave empty)
    start_blank = 3
    end_blank = start_blank + blank_rows - 1
    for r in range(start_blank, end_blank + 1):
        for c in range(cols):
            table.cell(r, c).text = ""

    # Total hours row at bottom: merge first three columns (Date + Time From + Time To)
    total_r = rows - 1
    table.cell(total_r, 0).merge(table.cell(total_r, 2))
    _set_paragraph_text(table.cell(total_r, 0), "Total hours", bold=True, size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Thicker outer border
    thick = {"sz": 12, "val": "single", "color": "000000"}
    for c in range(cols):
        _set_cell_border(table.cell(0, c), top=thick)
        _set_cell_border(table.cell(rows - 1, c), bottom=thick)
    for r in range(rows):
        _set_cell_border(table.cell(r, 0), left=thick)
        _set_cell_border(table.cell(r, cols - 1), right=thick)

    # Slightly thicker border around title row (top + bottom)
    for c in range(cols):
        _set_cell_border(table.cell(0, c), bottom=thick)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


if __name__ == "__main__":
    out = Path(__file__).resolve().parents[1] / "Log_Sheet_Work_Performed_During_Internship.docx"
    generate(out)
    print(f"Wrote: {out}")
